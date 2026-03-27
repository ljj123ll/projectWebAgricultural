#!/usr/bin/env python3
"""Extract text, tables, and OCR content from DOCX, PDF, and image files."""

from __future__ import annotations

import argparse
from pathlib import Path
import sys
import zipfile

import cv2
import docx
import fitz
import numpy as np
import pdfplumber
from pypdf import PdfReader
from rapidocr_onnxruntime import RapidOCR
import tabula


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


def format_markdown_table(rows: list[list[str]]) -> str:
    normalized = [[(cell or "").strip() for cell in row] for row in rows if row]
    normalized = [row for row in normalized if any(row)]
    if not normalized:
        return ""

    width = max(len(row) for row in normalized)
    padded = [row + [""] * (width - len(row)) for row in normalized]
    header = padded[0]
    separator = ["---"] * width
    body = padded[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def get_ocr_engine() -> RapidOCR:
    if not hasattr(get_ocr_engine, "_engine"):
        get_ocr_engine._engine = RapidOCR()
    return get_ocr_engine._engine


def ocr_image_array(image: np.ndarray) -> str:
    ocr = get_ocr_engine()
    result, _ = ocr(image)
    if not result:
        return ""
    return "\n".join(item[1].strip() for item in result if len(item) >= 2 and item[1].strip())


def ocr_image_path(path: Path) -> str:
    image = cv2.imread(str(path))
    if image is None:
        return ""
    return ocr_image_array(image)


def ocr_image_bytes(payload: bytes) -> str:
    image = cv2.imdecode(np.frombuffer(payload, dtype=np.uint8), cv2.IMREAD_COLOR)
    if image is None:
        return ""
    return ocr_image_array(image)


def read_docx(path: Path, ocr_images: bool = False) -> str:
    document = docx.Document(path)
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    if ocr_images:
        chunks.extend(extract_docx_image_ocr(path))

    return "\n".join(chunks).strip()


def extract_docx_image_ocr(path: Path) -> list[str]:
    chunks: list[str] = []
    with zipfile.ZipFile(path) as archive:
        image_names = [
            name for name in archive.namelist() if name.startswith("word/media/")
        ]
        for image_index, image_name in enumerate(image_names, start=1):
            text = ocr_image_bytes(archive.read(image_name)).strip()
            if text:
                chunks.append(f"[Image {image_index} OCR]")
                chunks.append(text)
    return chunks


def read_pdf_with_fitz(path: Path) -> str:
    chunks: list[str] = []
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if text:
                chunks.append(f"[Page {page_number}]")
                chunks.append(text)
    return "\n\n".join(chunks).strip()


def read_pdf_with_pdfplumber(path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                chunks.append(f"[Page {page_number}]")
                chunks.append(text)
    return "\n\n".join(chunks).strip()


def read_pdf_with_pypdf(path: Path) -> str:
    chunks: list[str] = []
    reader = PdfReader(path)
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append(f"[Page {page_number}]")
            chunks.append(text)
    return "\n\n".join(chunks).strip()


def extract_pdf_tables_with_pdfplumber(path: Path) -> list[str]:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        table_index = 1
        for page_number, page in enumerate(pdf.pages, start=1):
            for table in page.extract_tables() or []:
                markdown = format_markdown_table(
                    [[cell or "" for cell in row] for row in table if row]
                )
                if markdown:
                    chunks.append(f"[Table {table_index} from Page {page_number}]")
                    chunks.append(markdown)
                    table_index += 1
    return chunks


def extract_pdf_tables_with_tabula(path: Path) -> list[str]:
    chunks: list[str] = []
    table_index = 1
    for lattice in (False, True):
        try:
            tables = tabula.read_pdf(
                path,
                pages="all",
                multiple_tables=True,
                guess=not lattice,
                stream=not lattice,
                lattice=lattice,
                silent=True,
            )
        except Exception:
            continue

        for dataframe in tables or []:
            if dataframe.empty:
                continue
            rows = [list(map(str, dataframe.columns.tolist()))]
            rows.extend(
                [list(map(lambda value: "" if value is None else str(value), row)) for row in dataframe.fillna("").values.tolist()]
            )
            markdown = format_markdown_table(rows)
            if markdown:
                chunks.append(f"[Table {table_index} from Tabula]")
                chunks.append(markdown)
                table_index += 1
        if chunks:
            break
    return chunks


def ocr_pdf_pages(path: Path) -> list[str]:
    chunks: list[str] = []
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            pix = page.get_pixmap(dpi=200)
            image = cv2.imdecode(
                np.frombuffer(pix.tobytes("png"), dtype=np.uint8), cv2.IMREAD_COLOR
            )
            if image is None:
                continue
            text = ocr_image_array(image).strip()
            if text:
                chunks.append(f"[Page {page_number} OCR]")
                chunks.append(text)
    return chunks


def read_pdf(path: Path, force_ocr: bool = False, include_tables: bool = True) -> str:
    errors: list[str] = []
    chunks: list[str] = []

    for reader_name, reader_func in (
        ("PyMuPDF", read_pdf_with_fitz),
        ("pdfplumber", read_pdf_with_pdfplumber),
        ("pypdf", read_pdf_with_pypdf),
    ):
        try:
            text = reader_func(path)
            if text:
                chunks.append(text)
                break
            errors.append(f"{reader_name}: extracted empty text")
        except Exception as exc:  # pragma: no cover - fallback chain
            errors.append(f"{reader_name}: {exc}")

    if include_tables:
        table_chunks = extract_pdf_tables_with_pdfplumber(path)
        if not table_chunks:
            table_chunks = extract_pdf_tables_with_tabula(path)
        chunks.extend(table_chunks)

    if force_ocr or not chunks:
        chunks.extend(ocr_pdf_pages(path))

    if chunks:
        return "\n\n".join(chunk for chunk in chunks if chunk.strip()).strip()

    raise RuntimeError("Unable to extract text from PDF. " + "; ".join(errors))


def read_image(path: Path) -> str:
    text = ocr_image_path(path).strip()
    if not text:
        raise RuntimeError("No readable text was extracted from the image.")
    return text


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Read DOCX, PDF, or image files and print extracted text."
    )
    parser.add_argument("file", help="Path to a .docx, .pdf, or image file")
    parser.add_argument(
        "-o",
        "--output",
        help="Optional output text file path. If omitted, prints to stdout.",
    )
    parser.add_argument(
        "--ocr-images",
        action="store_true",
        help="OCR images embedded in DOCX files.",
    )
    parser.add_argument(
        "--ocr-pdf",
        action="store_true",
        help="Force OCR for all PDF pages, useful for scanned PDFs.",
    )
    parser.add_argument(
        "--no-tables",
        action="store_true",
        help="Skip PDF table extraction.",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    path = Path(args.file).expanduser().resolve()

    if not path.is_file():
        print(f"File not found: {path}", file=sys.stderr)
        return 1

    suffix = path.suffix.lower()
    if suffix == ".docx":
        text = read_docx(path, ocr_images=args.ocr_images)
    elif suffix == ".pdf":
        text = read_pdf(
            path,
            force_ocr=args.ocr_pdf,
            include_tables=not args.no_tables,
        )
    elif suffix in IMAGE_SUFFIXES:
        text = read_image(path)
    else:
        print("Only .docx, .pdf, and common image files are supported.", file=sys.stderr)
        return 1

    if not text:
        print("No readable text was extracted from the document.", file=sys.stderr)
        return 2

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
        output_path.write_text(text, encoding="utf-8")
    else:
        print(text)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
